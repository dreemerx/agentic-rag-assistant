"""
Document Loader Abstraction.

Design rationale:
- BaseLoader defines a single `load()` interface that returns a list of Document objects.
- Each file format has its own loader class — no if/elif dispatch.
- New formats are added by subclassing BaseLoader and registering in LOADER_REGISTRY.
- The upload endpoint uses LOADER_REGISTRY to pick the right loader by file extension.
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
import aiofiles


@dataclass
class Document:
    """A parsed document chunk with metadata for source citation."""

    content: str
    metadata: dict = field(default_factory=dict)


class BaseLoader(ABC):
    """Base class for all document loaders."""

    def __init__(self, file_path: str | Path):
        self.file_path = Path(file_path)

    @abstractmethod
    async def load(self) -> list[Document]:
        """Parse the file and return a list of Documents."""
        ...


class TXTLoader(BaseLoader):
    """Plain text loader."""

    async def load(self) -> list[Document]:
        async with aiofiles.open(self.file_path, "r", encoding="utf-8") as f:
            content = await f.read()
        return [Document(
            content=content,
            metadata={"source": str(self.file_path), "type": "txt"},
        )]


class MarkdownLoader(BaseLoader):
    """Markdown loader — reads raw markdown content."""

    async def load(self) -> list[Document]:
        async with aiofiles.open(self.file_path, "r", encoding="utf-8") as f:
            content = await f.read()
        return [Document(
            content=content,
            metadata={"source": str(self.file_path), "type": "markdown"},
        )]


class PDFLoader(BaseLoader):
    """PDF loader using pypdf for text extraction."""

    async def load(self) -> list[Document]:
        from pypdf import PdfReader

        reader = PdfReader(str(self.file_path))
        documents = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                documents.append(Document(
                    content=text.strip(),
                    metadata={
                        "source": str(self.file_path),
                        "type": "pdf",
                        "page": i + 1,
                    },
                ))
        return documents


class DOCXLoader(BaseLoader):
    """DOCX loader using python-docx."""

    async def load(self) -> list[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(self.file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return [Document(
            content=content,
            metadata={"source": str(self.file_path), "type": "docx"},
        )]


# === Loader Registry ===
# Maps file extensions to loader classes.
# Add new formats here — no other code changes needed.

LOADER_REGISTRY: dict[str, type[BaseLoader]] = {
    ".txt": TXTLoader,
    ".md": MarkdownLoader,
    ".markdown": MarkdownLoader,
    ".pdf": PDFLoader,
    ".docx": DOCXLoader,
}


def get_loader(file_path: str | Path) -> BaseLoader:
    """Get the appropriate loader for a file based on its extension."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    loader_cls = LOADER_REGISTRY.get(suffix)
    if not loader_cls:
        supported = list(LOADER_REGISTRY.keys())
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {supported}")
    return loader_cls(path)
